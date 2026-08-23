"""Turning an uploaded manuscript into text.

The extracted text is handed back to the student and put in front of them
before anything analyses it. That is not a convenience, it is what keeps the
rest of the system honest.

Every finding in a Weakness Map must quote the draft verbatim, and every quote
is verified against the text we were given. If extraction happened invisibly,
those quotes would be checked against a version of the manuscript the student
has never seen: a PDF with a two column layout, a running header repeated on
every page, a ligature that arrived as a different character. The student would
be shown a quote they cannot find in their own document, and told an examiner
will attack it.

So extraction produces text the student reads, edits if it is wrong, and
submits deliberately. What the analyzer sees is exactly what they saw.

Nothing is persisted here. The file is read in memory and discarded, and only
the text the student chose to submit goes any further.
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_PDF_PAGES = 400

# Below this, a PDF almost certainly has no text layer: it is a scan, or a set
# of page images. The distinction matters because the failure the student sees
# should be "this document has no selectable text", not "your draft is too
# short", which would send them looking for a problem in their writing.
MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER = 25


class ExtractionError(ValueError):
    """Any reason a document could not be turned into text."""


@dataclass
class ExtractedDraft:
    text: str
    page_count: int = 0
    notes: list[str] = field(default_factory=list)


def _clean(text: str) -> str:
    """Normalise extraction artefacts without touching the author's words.

    Ligatures and non-breaking spaces come out of PDFs as characters a student
    cannot type, so a quote containing one would never match a search in their
    own editor. Line endings are unified and runs of blank lines collapsed;
    paragraph breaks survive because they carry the structure the analyzer uses
    to tell one section from another.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace(" ", " ").replace("​", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # A hyphen at a line break is a typesetting artefact, not spelling.
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _drop_repeated_lines(pages: list[str]) -> tuple[list[str], list[str]]:
    """Remove headers and footers that repeat on most pages.

    A running header quoted back as a weakness would be absurd, and page numbers
    scattered through the text break sentences apart where the analyzer looks
    for them. Only short lines that appear on more than half the pages are
    removed, so a repeated sentence of actual argument survives.
    """
    if len(pages) < 3:
        return pages, []

    counts: dict[str, int] = {}
    for page in pages:
        for line in {ln.strip() for ln in page.split("\n") if ln.strip()}:
            if len(line) <= 80:
                counts[line] = counts.get(line, 0) + 1

    threshold = len(pages) // 2 + 1
    repeated = {line for line, count in counts.items() if count >= threshold}
    if not repeated:
        return pages, []

    cleaned = [
        "\n".join(ln for ln in page.split("\n") if ln.strip() not in repeated) for page in pages
    ]
    notes = [
        f"Removed {len(repeated)} line(s) repeating on most pages, such as a running "
        f"header or page number."
    ]
    return cleaned, notes


def extract_pdf(data: bytes) -> ExtractedDraft:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise ExtractionError(f"This PDF could not be read: {exc}") from exc

    if reader.is_encrypted:
        # An empty password unlocks many "protected" PDFs. Anything else is a
        # password we do not have and will not ask for.
        try:
            if reader.decrypt("") == 0:
                raise ExtractionError(
                    "This PDF is password protected. Remove the password and upload it again."
                )
        except ExtractionError:
            raise
        except Exception as exc:  # noqa: BLE001 - any failure here means locked
            raise ExtractionError(
                "This PDF is password protected. Remove the password and upload it again."
            ) from exc

    page_count = len(reader.pages)
    if page_count == 0:
        raise ExtractionError("This PDF has no pages.")
    if page_count > MAX_PDF_PAGES:
        raise ExtractionError(
            f"This PDF has {page_count} pages, and {MAX_PDF_PAGES} is the limit. "
            "Upload the chapters you want examined."
        )

    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page must not lose the rest
            logger.warning("Page %d of the PDF could not be read; skipping it.", index)
            pages.append("")

    pages, notes = _drop_repeated_lines(pages)
    text = _clean("\n\n".join(pages))

    if len(text) < page_count * MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER:
        raise ExtractionError(
            "This PDF has no selectable text. It is most likely a scan, so the "
            "words are pictures rather than characters. Upload the original "
            "document, or paste the text directly."
        )

    return ExtractedDraft(text=text, page_count=page_count, notes=notes)


def extract_docx(data: bytes) -> ExtractedDraft:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - see below
        # Deliberately broad. A .docx is a zip of XML, and a file that is not
        # one fails somewhere down that stack: BadZipFile from the standard
        # library, KeyError on a missing part, PackageNotFoundError from
        # python-docx. Every one of them means the same thing to the student,
        # and none of them should reach them as a stack trace about zip files.
        raise ExtractionError(
            "This file is not a valid .docx. If it is an older .doc, save it as .docx first."
        ) from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]

    # Tables carry methodology and results in most theses, and dropping them
    # would silently remove the sections an examiner presses hardest on.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = _clean("\n".join(blocks))
    if not text:
        raise ExtractionError("This document contains no text.")

    return ExtractedDraft(text=text)


def extract_draft(filename: str, data: bytes) -> ExtractedDraft:
    """Extract text from an uploaded manuscript.

    The format is decided by the file extension rather than by sniffing the
    bytes, because the student chose the file and an honest error naming their
    file is more useful than a clever guess about its contents.
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ExtractionError(
            f"This file is {len(data) // (1024 * 1024)} MB, and "
            f"{MAX_FILE_BYTES // (1024 * 1024)} MB is the limit."
        )

    lowered = filename.lower().strip()
    if lowered.endswith(".pdf"):
        return extract_pdf(data)
    if lowered.endswith(".docx"):
        return extract_docx(data)
    if lowered.endswith(".txt") or lowered.endswith(".md"):
        return ExtractedDraft(text=_clean(data.decode("utf-8", errors="replace")))

    raise ExtractionError(
        "Only PDF, DOCX, and plain text files can be read. If your draft is a "
        "Google Doc, download it as .docx first."
    )
