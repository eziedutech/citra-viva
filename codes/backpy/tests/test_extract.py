"""Manuscript extraction tests.

Documents are built in the test rather than committed as fixtures, so what each
test exercises is visible in the test itself: a scanned page is a page with no
text on it, a running header is a line repeated across pages, and neither is a
binary blob a reader has to take on trust.

The failures matter as much as the successes here. A student who uploads a
scanned thesis needs to be told their document has no selectable text, not that
their draft is too short, which would send them looking for a problem in their
writing that does not exist.
"""

from __future__ import annotations

import io

import pytest

from app.ingest.extract import (
    MAX_FILE_BYTES,
    ExtractionError,
    extract_docx,
    extract_draft,
    extract_pdf,
)

BODY = (
    "Penelitian ini menggunakan desain survei potong lintang terhadap 120 mahasiswa "
    "Fakultas Ekonomi yang dipilih menggunakan teknik accidental sampling."
)


def build_pdf(pages: list[list[str]], *, encrypt: str | None = None) -> bytes:
    """A PDF with real text on it, one list of lines per page."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    if encrypt is not None:
        pdf.setEncrypt(encrypt)
    for lines in pages:
        y = 800
        for line in lines:
            pdf.drawString(60, y, line)
            y -= 18
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def build_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                added.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


def test_a_text_pdf_comes_back_as_text():
    result = extract_pdf(build_pdf([[BODY], ["Hasil menunjukkan korelasi negatif."]]))

    assert result.page_count == 2
    assert "accidental sampling" in result.text
    assert "korelasi negatif" in result.text


def test_a_scanned_pdf_says_it_has_no_text_rather_than_looking_short():
    """A blank page stands in for a scan: pixels, no characters."""
    with pytest.raises(ExtractionError, match="no selectable text"):
        extract_pdf(build_pdf([[], [], []]))


def test_a_password_protected_pdf_is_named_as_such():
    with pytest.raises(ExtractionError, match="password protected"):
        extract_pdf(build_pdf([[BODY]], encrypt="rahasia"))


def test_a_file_that_is_not_a_pdf_is_refused_clearly():
    with pytest.raises(ExtractionError):
        extract_pdf(b"this is not a pdf at all")


def test_running_headers_are_removed_but_the_argument_is_not():
    """A header quoted back as a weakness would be absurd."""
    header = "Universitas Nusantara"
    pages = [[header, BODY], [header, "Bab dua."], [header, "Bab tiga."], [header, "Bab empat."]]

    result = extract_pdf(build_pdf(pages))

    assert header not in result.text
    assert "accidental sampling" in result.text
    assert result.notes


def test_a_repeated_line_across_only_two_pages_survives():
    """Two pages is not evidence of a running header, and removing a real
    sentence would silently edit the student's manuscript."""
    line = "Temuan ini konsisten pada kedua kelompok."
    result = extract_pdf(build_pdf([[line, BODY], [line, "Lanjutan."]]))

    assert line in result.text


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


def test_a_docx_comes_back_as_text():
    result = extract_docx(build_docx(["BAB I PENDAHULUAN", BODY]))

    assert "BAB I PENDAHULUAN" in result.text
    assert "accidental sampling" in result.text


def test_table_contents_are_kept():
    """Methodology and results live in tables in most theses, and those are the
    sections an examiner presses hardest on."""
    result = extract_docx(
        build_docx(
            ["Hasil analisis:"],
            table=[["Variabel", "Koefisien"], ["Durasi media sosial", "-0,42"]],
        )
    )

    assert "Koefisien" in result.text
    assert "-0,42" in result.text


def test_an_empty_document_is_refused():
    with pytest.raises(ExtractionError, match="no text"):
        extract_docx(build_docx([]))


def test_a_doc_that_is_not_a_docx_is_refused_with_advice():
    with pytest.raises(ExtractionError, match="older .doc"):
        extract_docx(b"\xd0\xcf\x11\xe0 old binary word format")


# --------------------------------------------------------------------------- #
# Dispatch and limits
# --------------------------------------------------------------------------- #


def test_the_extension_decides_the_format():
    assert "accidental sampling" in extract_draft("bab3.pdf", build_pdf([[BODY]])).text
    assert "accidental sampling" in extract_draft("bab3.docx", build_docx([BODY])).text
    assert "accidental sampling" in extract_draft("bab3.txt", BODY.encode("utf-8")).text


def test_an_unsupported_format_says_what_to_do_instead():
    with pytest.raises(ExtractionError, match="Google Doc"):
        extract_draft("draft.pages", b"anything")


def test_an_empty_upload_is_refused():
    with pytest.raises(ExtractionError, match="empty"):
        extract_draft("draft.pdf", b"")


def test_an_oversized_upload_is_refused_before_it_is_parsed():
    with pytest.raises(ExtractionError, match="MB is the limit"):
        extract_draft("draft.pdf", b"x" * (MAX_FILE_BYTES + 1))


def test_extraction_artefacts_are_normalised():
    """A ligature or a hyphenated line break produces characters a student
    cannot type, so a quote containing one would never match a search in their
    own editor."""
    result = extract_draft(
        "draft.txt",
        "Peneliti mengiden-\ntifikasi konﬂik kepentingan.".encode(),
    )

    assert "mengidentifikasi" in result.text
    assert "konflik kepentingan" in result.text
